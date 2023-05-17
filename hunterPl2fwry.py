

# 加载其他模块和库
import openpyxl
import re,requests
from docx import Document
from concurrent.futures import ThreadPoolExecutor
from argparse import ArgumentParser
from fanweiry import scan_bug1
import fanweiaes
import os

parse = ArgumentParser(description="jwf")
parse.add_argument('-f', "--file", dest="file", type=str)
parse.add_argument('-id', "--id", dest="id", type=str)  # 这里添加新的参数
args = parse.parse_args()

input_file = args.file.replace('.\\', '')  # 获取要处理的文件名
id = args.id  # 获取输入的payload
target = str(input_file)
exec = ThreadPoolExecutor(max_workers=200)



def create_word_file(gunit_name, s,ip,port):
    # 创建一个新的 Word 文档
    doc = Document()
    # 添加内容
    doc.add_paragraph(f'{gunit_name}存在泛微任意用户登录')
    doc.add_paragraph(f'访问下面链接即可登录用户')
    doc.add_paragraph(f'{s}')

    # 保存文档
    pattern = r"[/<>\s]+"
    replaced_text = re.sub(pattern, "", gunit_name)
    doc.save(f'output/9{replaced_text}{ip}存在泛微任意用户登录.docx')
    with open("res.txt","a+") as f:
        f.write(f'output/9{replaced_text}存在泛微任意用户登录.docx')


def start1(url,GSName,ip,port):
    # print(url)
    code,url1 = scan_bug1(url)

    if(code==1):
        print(url1)
        create_word_file(GSName,url1,ip,port)

# 加载工作簿
workbook = openpyxl.load_workbook(input_file)
# 获取第一个工作表
worksheet = workbook.active

# 要提取的列
column_a_index = 0  # A列的索引
column_n_index = 13 # N列的索引
column_ip_index = 2
column_port_index = 4
# 遍历每一行
for row in worksheet.iter_rows(min_row=1, max_row=worksheet.max_row):
    # 获取第A列和第N列的单元格
    cell_a = row[column_a_index]
    cell_n = row[column_n_index]
    cell_ip = row[column_ip_index]
    cell_port = row[column_port_index]
    # 输出值
    exec.submit(start1,cell_a.value,cell_n.value,cell_ip.value,cell_port.value)

# 关闭工作簿
workbook.close()
# 等待 Ctrl+C 信号

